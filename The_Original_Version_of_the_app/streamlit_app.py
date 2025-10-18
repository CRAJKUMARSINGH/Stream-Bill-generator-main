import streamlit as st
import pandas as pd
import pdfkit
from docx import Document
from docx.shared import Pt
from num2words import num2words
import os
import zipfile
import tempfile
from jinja2 import Environment, FileSystemLoader
from pypdf import PdfReader, PdfWriter
import numpy as np
import platform
from datetime import datetime
import traceback

# Set up Jinja2 environment
env = Environment(loader=FileSystemLoader("templates"), cache_size=0)

# Temporary directory
TEMP_DIR = tempfile.mkdtemp()

# Configure wkhtmltopdf
if platform.system() == "Windows":
    wkhtmltopdf_path = r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe"
    config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
else:
    config = pdfkit.configuration()

def number_to_words(number):
    try:
        return num2words(int(number), lang="en_IN").title()
    except:
        return str(number)
##########################################################################################
def process_bill(ws_wo, ws_bq, ws_extra, premium_percent, premium_type):
    st.write("Starting process_bill")
    first_page_data = {"header": [], "items": [], "totals": {}}
    last_page_data = {"payable_amount": 0, "amount_words": ""}
    deviation_data = {"items": [], "summary": {}}
    extra_items_data = {"items": []}
    note_sheet_data = {"notes": []}
################################################################################################################
    # Header (A1:I19)
    #header_data = ws_wo.iloc[:19].replace(np.nan, "").values.tolist()
    #first_page_data["header"] = header_data
    ###### REPLACEMENT 18 APRIL 2025
    from datetime import datetime, date

    # Header (A1:G19) only — matching actual data range
    header_data = ws_wo.iloc[:19, :7].replace(np.nan, "").values.tolist()

    # Ensure all dates are formatted as date-only strings (optional step, if needed before saving)
    for i in range(len(header_data)):
        for j in range(len(header_data[i])):
            val = header_data[i][j]
            if isinstance(val, (pd.Timestamp, datetime, date)):
                header_data[i][j] = val.strftime("%d-%m-%Y")

    # Assign to first page
    first_page_data["header"] = header_data
############################################################################################################
    # Work Order items
    last_row_wo = ws_wo.shape[0]
    for i in range(21, last_row_wo):
        qty_raw = ws_bq.iloc[i, 3] if i < ws_bq.shape[0] and pd.notnull(ws_bq.iloc[i, 3]) else None
        rate_raw = ws_wo.iloc[i, 4] if pd.notnull(ws_wo.iloc[i, 4]) else None

        qty = 0
        if isinstance(qty_raw, (int, float)):
            qty = float(qty_raw)
        elif isinstance(qty_raw, str):
            cleaned_qty = qty_raw.strip().replace(',', '').replace(' ', '')
            try:
                qty = float(cleaned_qty)
            except ValueError:
                st.warning(f"Skipping invalid quantity at Bill Quantity row {i+1}: '{qty_raw}'")
                continue

        rate = 0
        if isinstance(rate_raw, (int, float)):
            rate = float(rate_raw)
        elif isinstance(rate_raw, str):
            cleaned_rate = rate_raw.strip().replace(',', '').replace(' ', '')
            try:
                rate = float(cleaned_rate)
            except ValueError:
                st.warning(f"Skipping invalid rate at Work Order row {i+1}: '{rate_raw}'")
                continue

        item = {
            "serial_no": str(ws_wo.iloc[i, 0]) if pd.notnull(ws_wo.iloc[i, 0]) else "",
            "description": str(ws_wo.iloc[i, 1]) if pd.notnull(ws_wo.iloc[i, 1]) else "",
            "unit": str(ws_wo.iloc[i, 2]) if pd.notnull(ws_wo.iloc[i, 2]) else "",
            "quantity": qty,
            "rate": rate,
            "remark": str(ws_wo.iloc[i, 6]) if pd.notnull(ws_wo.iloc[i, 6]) else "",
            "amount": round(qty * rate) if qty and rate else 0,
            "is_divider": False
        }
        first_page_data["items"].append(item)

    # Extra Items divider
    first_page_data["items"].append({
        "description": "Extra Items (With Premium)",
        "bold": True,
        "underline": True,
        "amount": 0,
        "quantity": 0,
        "rate": 0,
        "serial_no": "",
        "unit": "",
        "remark": "",
        "is_divider": True
    })

    # Extra Items
    last_row_extra = ws_extra.shape[0]
    for j in range(6, last_row_extra):
        qty_raw = ws_extra.iloc[j, 3] if pd.notnull(ws_extra.iloc[j, 3]) else None
        rate_raw = ws_extra.iloc[j, 5] if pd.notnull(ws_extra.iloc[j, 5]) else None

        qty = 0
        if isinstance(qty_raw, (int, float)):
            qty = float(qty_raw)
        elif isinstance(qty_raw, str):
            cleaned_qty = qty_raw.strip().replace(',', '').replace(' ', '')
            try:
                qty = float(cleaned_qty)
            except ValueError:
                st.warning(f"Skipping invalid quantity at Extra Items row {j+1}: '{qty_raw}'")
                continue

        rate = 0
        if isinstance(rate_raw, (int, float)):
            rate = float(rate_raw)
        elif isinstance(rate_raw, str):
            cleaned_rate = rate_raw.strip().replace(',', '').replace(' ', '')
            try:
                rate = float(cleaned_rate)
            except ValueError:
                st.warning(f"Skipping invalid rate at Extra Items row {j+1}: '{rate_raw}'")
                continue

        item = {
            "serial_no": str(ws_extra.iloc[j, 0]) if pd.notnull(ws_extra.iloc[j, 0]) else "",
            "description": str(ws_extra.iloc[j, 2]) if pd.notnull(ws_extra.iloc[j, 2]) else "",
            "unit": str(ws_extra.iloc[j, 4]) if pd.notnull(ws_extra.iloc[j, 4]) else "",
            "quantity": qty,
            "rate": rate,
            "remark": str(ws_extra.iloc[j, 1]) if pd.notnull(ws_extra.iloc[j, 1]) else "",
            "amount": round(qty * rate) if qty and rate else 0,
            "is_divider": False
        }
        first_page_data["items"].append(item)
        extra_items_data["items"].append(item.copy())  # Copy for standalone Extra Items

    # Totals
    data_items = [item for item in first_page_data["items"] if not item.get("is_divider", False)]
    total_amount = round(sum(item.get("amount", 0) for item in data_items))
    premium_amount = round(total_amount * (premium_percent / 100) if premium_type == "above" else -total_amount * (premium_percent / 100))
    payable_amount = round(total_amount + premium_amount)

    first_page_data["totals"] = {
        "grand_total": total_amount,
        "premium": {"percent": premium_percent / 100, "type": premium_type, "amount": premium_amount},
        "payable": payable_amount
    }

    try:
        extra_items_start = next(i for i, item in enumerate(first_page_data["items"]) if item.get("description") == "Extra Items (With Premium)")
        extra_items = [item for item in first_page_data["items"][extra_items_start + 1:] if not item.get("is_divider", False)]
        extra_items_sum = round(sum(item.get("amount", 0) for item in extra_items))
        extra_items_premium = round(extra_items_sum * (premium_percent / 100) if premium_type == "above" else -extra_items_sum * (premium_percent / 100))
        first_page_data["totals"]["extra_items_sum"] = extra_items_sum + extra_items_premium
    except StopIteration:
        first_page_data["totals"]["extra_items_sum"] = 0

    # Last Page
    last_page_data = {"payable_amount": payable_amount, "amount_words": number_to_words(payable_amount)}

    # Deviation Statement
    work_order_total = 0
    executed_total = 0
    overall_excess = 0
    overall_saving = 0
    for i in range(21, last_row_wo):
        st.write(f"Processing deviation row {i+1}: wo_qty={ws_wo.iloc[i, 3]}, wo_rate={ws_wo.iloc[i, 4]}, bq_qty={ws_bq.iloc[i, 3] if i < ws_bq.shape[0] else 'N/A'}")
        qty_wo_raw = ws_wo.iloc[i, 3] if pd.notnull(ws_wo.iloc[i, 3]) else None
        rate_raw = ws_wo.iloc[i, 4] if pd.notnull(ws_wo.iloc[i, 4]) else None
        qty_bill_raw = ws_bq.iloc[i, 3] if i < ws_bq.shape[0] and pd.notnull(ws_bq.iloc[i, 3]) else None

        qty_wo = 0
        if isinstance(qty_wo_raw, (int, float)):
            qty_wo = float(qty_wo_raw)
        elif isinstance(qty_wo_raw, str):
            cleaned_qty_wo = qty_wo_raw.strip().replace(',', '').replace(' ', '')
            try:
                qty_wo = float(cleaned_qty_wo)
            except ValueError:
                st.warning(f"Skipping invalid qty_wo at row {i+1}: '{qty_wo_raw}'")
                continue

        rate = 0
        if isinstance(rate_raw, (int, float)):
            rate = float(rate_raw)
        elif isinstance(rate_raw, str):
            cleaned_rate = rate_raw.strip().replace(',', '').replace(' ', '')
            try:
                rate = float(cleaned_rate)
            except ValueError:
                st.warning(f"Skipping invalid rate at row {i+1}: '{rate_raw}'")
                continue

        qty_bill = 0
        if isinstance(qty_bill_raw, (int, float)):
            qty_bill = float(qty_bill_raw)
        elif isinstance(qty_bill_raw, str):
            cleaned_qty_bill = qty_bill_raw.strip().replace(',', '').replace(' ', '')
            try:
                qty_bill = float(cleaned_qty_bill)
            except ValueError:
                st.warning(f"Skipping invalid qty_bill at row {i+1}: '{qty_bill_raw}'")
                continue

        amt_wo = round(qty_wo * rate)
        amt_bill = round(qty_bill * rate)
        excess_qty = qty_bill - qty_wo if qty_bill > qty_wo else 0
        excess_amt = round(excess_qty * rate) if excess_qty > 0 else 0
        saving_qty = qty_wo - qty_bill if qty_bill < qty_wo else 0
        saving_amt = round(saving_qty * rate) if saving_qty > 0 else 0

        item = {
            "serial_no": str(ws_wo.iloc[i, 0]) if pd.notnull(ws_wo.iloc[i, 0]) else "",
            "description": str(ws_wo.iloc[i, 1]) if pd.notnull(ws_wo.iloc[i, 1]) else "",
            "unit": str(ws_wo.iloc[i, 2]) if pd.notnull(ws_wo.iloc[i, 2]) else "",
            "qty_wo": qty_wo,
            "rate": rate,
            "amt_wo": amt_wo,
            "qty_bill": qty_bill,
            "amt_bill": amt_bill,
            "excess_qty": excess_qty,
            "excess_amt": excess_amt,
            "saving_qty": saving_qty,
            "saving_amt": saving_amt
        }
        deviation_data["items"].append(item)
        work_order_total += amt_wo
        executed_total += amt_bill
        overall_excess += excess_amt
        overall_saving += saving_amt

    # Deviation Summary
    tender_premium_f = round(work_order_total * (premium_percent / 100) if premium_type == "above" else -work_order_total * (premium_percent / 100))
    tender_premium_h = round(executed_total * (premium_percent / 100) if premium_type == "above" else -executed_total * (premium_percent / 100))
    tender_premium_j = round(overall_excess * (premium_percent / 100) if premium_type == "above" else -overall_excess * (premium_percent / 100))
    tender_premium_l = round(overall_saving * (premium_percent / 100) if premium_type == "above" else -overall_saving * (premium_percent / 100))
    grand_total_f = work_order_total + tender_premium_f
    grand_total_h = executed_total + tender_premium_h
    grand_total_j = overall_excess + tender_premium_j
    grand_total_l = overall_saving + tender_premium_l
    net_difference = grand_total_h - grand_total_f

    deviation_data["summary"] = {
        "work_order_total": round(work_order_total),
        "executed_total": round(executed_total),
        "overall_excess": round(overall_excess),
        "overall_saving": round(overall_saving),
        "premium": {"percent": premium_percent / 100, "type": premium_type},
        "tender_premium_f": tender_premium_f,
        "tender_premium_h": tender_premium_h,
        "tender_premium_j": tender_premium_j,
        "tender_premium_l": tender_premium_l,
        "grand_total_f": grand_total_f,
        "grand_total_h": grand_total_h,
        "grand_total_j": grand_total_j,
        "grand_total_l": grand_total_l,
        "net_difference": round(net_difference)
    }

    st.write(f"first_page_data['items'] type: {type(first_page_data['items'])}, length: {len(first_page_data['items'])}")
    st.write(f"extra_items_data['items'] type: {type(extra_items_data['items'])}, length: {len(extra_items_data['items'])}")
    st.write(f"deviation_data['items'] type: {type(deviation_data['items'])}, length: {len(deviation_data['items'])}")
    return first_page_data, last_page_data, deviation_data, extra_items_data, note_sheet_data
########################################################################################################################################################
def generate_bill_notes(payable_amount, work_order_amount, extra_item_amount):
    percentage_work_done = float(payable_amount / work_order_amount * 100) if work_order_amount > 0 else 0
    serial_number = 1
    note = []
    note.append(f"{serial_number}. The work has been completed {percentage_work_done:.2f}% of the Work Order Amount.")
    serial_number += 1
    if percentage_work_done < 90:
        note.append(f"{serial_number}. The execution of work at final stage is less than 90%...")
        serial_number += 1
    elif percentage_work_done > 100 and percentage_work_done <= 105:
        note.append(f"{serial_number}. Requisite Deviation Statement is enclosed...")
        serial_number += 1
    elif percentage_work_done > 105:
        note.append(f"{serial_number}. Requisite Deviation Statement is enclosed...")
        serial_number += 1
    note.append(f"{serial_number}. Quality Control (QC) test reports attached.")
    serial_number += 1
    if extra_item_amount > 0:
        extra_item_percentage = float(extra_item_amount / work_order_amount * 100) if work_order_amount > 0 else 0
        if extra_item_percentage > 5:
            note.append(f"{serial_number}. The amount of Extra items is Rs. {extra_item_amount}...")
        else:
            note.append(f"{serial_number}. The amount of Extra items is Rs. {extra_item_amount}...")
        serial_number += 1
    note.append(f"{serial_number}. Please peruse above details for necessary decision-making.")
    note.append("")
    note.append("                                Premlata Jain")
    note.append("                               AAO- As Auditor")
    return {"notes": note}

def generate_pdf(sheet_name, data, orientation, output_path):
    st.write(f"Generating PDF for {sheet_name}, data type: {type(data)}, items type: {type(data.get('items', []))}, totals.premium.percent: {data.get('totals', {}).get('premium', {}).get('percent', 'N/A')}")
    try:
        template = env.get_template(f"{sheet_name.lower().replace(' ', '_')}.html")
        html_content = template.render(data=data)
        options = {
            "page-size": "A4",
            "orientation": orientation,
        }
        # Apply margins only to Note Sheet
        if sheet_name != "Note Sheet":
            options.update({
                "margin-top": "0in",
                "margin-bottom": "0in",
                "margin-left": "0in",
                "margin-right": "0in"
            })
        else:
            options.update({
                "margin-top": "0.25in",
                "margin-bottom": "0.6in",
                "margin-left": "0.25in",
                "margin-right": "0.25in"
            })
        pdfkit.from_string(
            html_content,
            output_path,
            configuration=config,
            options=options
        )
        st.write(f"Finished PDF for {sheet_name}")
    except Exception as e:
        st.error(f"Error generating PDF for {sheet_name}: {str(e)}")
        st.write(traceback.format_exc())
        raise

def create_word_doc(sheet_name, data, doc_path):
    st.write(f"Creating Word doc for {sheet_name}")
    try:
        doc = Document()
        if sheet_name == "First Page":
            table = doc.add_table(rows=len(data["items"]) + 3, cols=9)
            table.style = "Table Grid"
            for i, item in enumerate(data["items"]):
                row = table.rows[i]
                row.cells[0].text = str(item.get("unit", ""))
                row.cells[2].text = str(item.get("quantity", ""))
                row.cells[3].text = str(item.get("serial_no", ""))
                row.cells[4].text = str(item.get("description", ""))
                row.cells[5].text = str(item.get("rate", ""))
                row.cells[6].text = str(item.get("amount", ""))
                row.cells[8].text = str(item.get("remark", ""))
            row = table.rows[-3]
            row.cells[4].text = "Grand Total"
            row.cells[6].text = str(data["totals"].get("grand_total", ""))
            row = table.rows[-2]
            row.cells[4].text = f"Tender Premium @ {data['totals']['premium'].get('percent', 0):.2%}"
            row.cells[6].text = str(data["totals"]["premium"].get("amount", ""))
            row = table.rows[-1]
            row.cells[4].text = "Payable Amount"
            row.cells[6].text = str(data["totals"].get("payable", ""))
        elif sheet_name == "Last Page":
            doc.add_paragraph(f"Payable Amount: {data.get('payable_amount', '')}")
            doc.add_paragraph(f"Total in Words: {data.get('amount_words', '')}")
        elif sheet_name == "Extra Items":
            table = doc.add_table(rows=len(data["items"]) + 1, cols=7)
            table.style = "Table Grid"
            headers = ["Serial No.", "Remark", "Description", "Quantity", "Unit", "Rate", "Amount"]
            for j, header in enumerate(headers):
                table.rows[0].cells[j].text = header
            for i, item in enumerate(data["items"]):
                row = table.rows[i + 1]
                row.cells[0].text = str(item.get("serial_no", ""))
                row.cells[1].text = str(item.get("remark", ""))
                row.cells[2].text = str(item.get("description", ""))
                row.cells[3].text = str(item.get("quantity", ""))
                row.cells[4].text = str(item.get("unit", ""))
                row.cells[5].text = str(item.get("rate", ""))
                row.cells[6].text = str(item.get("amount", ""))
        elif sheet_name == "Deviation Statement":
            table = doc.add_table(rows=len(data["items"]) + 5, cols=12)
            table.style = "Table Grid"
            headers = ["Serial No.", "Description", "Unit", "Qty WO", "Rate", "Amt WO", "Qty Bill", "Amt Bill", "Excess Qty", "Excess Amt", "Saving Qty", "Saving Amt"]
            for j, header in enumerate(headers):
                table.rows[0].cells[j].text = header
            for i, item in enumerate(data["items"]):
                row = table.rows[i + 1]
                row.cells[0].text = str(item.get("serial_no", ""))
                row.cells[1].text = str(item.get("description", ""))
                row.cells[2].text = str(item.get("unit", ""))
                row.cells[3].text = str(item.get("qty_wo", ""))
                row.cells[4].text = str(item.get("rate", ""))
                row.cells[5].text = str(item.get("amt_wo", ""))
                row.cells[6].text = str(item.get("qty_bill", ""))
                row.cells[7].text = str(item.get("amt_bill", ""))
                row.cells[8].text = str(item.get("excess_qty", ""))
                row.cells[9].text = str(item.get("excess_amt", ""))
                row.cells[10].text = str(item.get("saving_qty", ""))
                row.cells[11].text = str(item.get("saving_amt", ""))
            row = table.rows[-4]
            row.cells[1].text = "Grand Total"
            row.cells[5].text = str(data["summary"].get("work_order_total", ""))
            row.cells[7].text = str(data["summary"].get("executed_total", ""))
            row.cells[9].text = str(data["summary"].get("overall_excess", ""))
            row.cells[11].text = str(data["summary"].get("overall_saving", ""))
            row = table.rows[-3]
            row.cells[1].text = f"Add Tender Premium ({data['summary']['premium'].get('percent', 0):.2%})"
            row.cells[5].text = str(data["summary"].get("tender_premium_f", ""))
            row.cells[7].text = str(data["summary"].get("tender_premium_h", ""))
            row.cells[9].text = str(data["summary"].get("tender_premium_j", ""))
            row.cells[11].text = str(data["summary"].get("tender_premium_l", ""))
            row = table.rows[-2]
            row.cells[1].text = "Grand Total including Tender Premium"
            row.cells[5].text = str(data["summary"].get("grand_total_f", ""))
            row.cells[7].text = str(data["summary"].get("grand_total_h", ""))
            row.cells[9].text = str(data["summary"].get("grand_total_j", ""))
            row.cells[11].text = str(data["summary"].get("grand_total_l", ""))
            row = table.rows[-1]
            net_difference = data["summary"].get("net_difference", 0)
            row.cells[1].text = "Overall Excess" if net_difference > 0 else "Overall Saving"
            row.cells[7].text = str(abs(round(net_difference)))
        elif sheet_name == "Note Sheet":
            for note in data.get("notes", []):
                doc.add_paragraph(str(note))
        doc.save(doc_path)
        st.write(f"Finished Word doc for {sheet_name}")
    except Exception as e:
        st.error(f"Error creating Word doc for {sheet_name}: {str(e)}")
        raise

# Streamlit app
st.title("Bill Generator")
st.write("Upload an Excel file and enter tender premium details.")

uploaded_file = st.file_uploader("Choose an Excel file", type="xlsx")
premium_percent = st.number_input("Tender Premium %", min_value=0.0, max_value=100.0, step=0.01)
premium_type = st.selectbox("Premium Type", ["Above", "Below"])

if uploaded_file is not None and st.button("Generate Bill"):
    try:
        xl = pd.ExcelFile(uploaded_file)
        ws_wo = xl.parse("Work Order", header=None)
        ws_bq = xl.parse("Bill Quantity", header=None)
        ws_extra = xl.parse("Extra Items", header=None)

        first_page_data, last_page_data, deviation_data, extra_items_data, note_sheet_data = process_bill(
            ws_wo, ws_bq, ws_extra, premium_percent, premium_type.lower()
        )

                # Generate note sheet
        try:
            work_order_amount = sum(
                float(ws_wo.iloc[i, 3]) * float(ws_wo.iloc[i, 4])
                for i in range(21, ws_wo.shape[0])
                if pd.notnull(ws_wo.iloc[i, 3]) and pd.notnull(ws_wo.iloc[i, 4])
            )
        except Exception as e:
            st.error(f"Error calculating work_order_amount: {e}")
            work_order_amount = 854678  # Fallback value

        extra_item_amount = first_page_data["totals"].get("extra_items_sum", 0)
        payable_amount = first_page_data["totals"].get("payable", 0)
        note_sheet_data = generate_bill_notes(payable_amount, work_order_amount, extra_item_amount)

        # Define work_order_data from ws_wo or a Work Order sheet
        work_order_data = {
            'agreement_no': ws_wo.iloc[0, 1] if pd.notnull(ws_wo.iloc[0, 1]) else '48/2024-25',
            'name_of_work': ws_wo.iloc[1, 1] if pd.notnull(ws_wo.iloc[1, 1]) else 'Electric Repair and MTC work at Govt. Ambedkar hostel Ambamata, Govardhanvilas, Udaipur',
            'name_of_firm': ws_wo.iloc[2, 1] if pd.notnull(ws_wo.iloc[2, 1]) else 'M/s Seema Electrical Udaipur',
            'date_commencement': ws_wo.iloc[3, 1] if pd.notnull(ws_wo.iloc[3, 1]) else '18/01/2025',
            'date_completion': ws_wo.iloc[4, 1] if pd.notnull(ws_wo.iloc[4, 1]) else '17/04/2025',
            'actual_completion': ws_wo.iloc[5, 1] if pd.notnull(ws_wo.iloc[5, 1]) else '01/03/2025',
            'work_order_amount': str(work_order_amount)
        }

        # Prepare note_sheet_data with VBA-style notes
        percentage_work_done = (float(payable_amount) / float(work_order_amount) * 100) if work_order_amount > 0 else 0
        notes = [
            f"1. The work has been completed {percentage_work_done:.2f}% of the Work Order Amount."
        ]
        if percentage_work_done < 90:
            notes.append("2. The execution of work at final stage is less than 90% of the Work Order Amount, the Requisite Deviation Statement is enclosed to observe check on unuseful expenditure. Approval of the Deviation is having jurisdiction under this office.")
        elif 100 < percentage_work_done <= 105:
            notes.append("2. Requisite Deviation Statement is enclosed. The Overall Excess is less than or equal to 5% and is having approval jurisdiction under this office.")
        elif percentage_work_done > 105:
            notes.append("2. Requisite Deviation Statement is enclosed. The Overall Excess is more than 5% and Approval of the Deviation Case is required from the Superintending Engineer, PWD Electrical Circle, Udaipur.")
        delay_days = (datetime.strptime(work_order_data['actual_completion'], '%d/%m/%Y') - datetime.strptime(work_order_data['date_completion'], '%d/%m/%Y')).days
        if delay_days > 0:
            time_allowed = (datetime.strptime(work_order_data['date_completion'], '%d/%m/%Y') - datetime.strptime(work_order_data['date_commencement'], '%d/%m/%Y')).days
            notes.append(f"3. Time allowed for completion of the work was {time_allowed} days. The Work was delayed by {delay_days} days.")
            if delay_days > 0.5 * time_allowed:
                notes.append("4. Approval of the Time Extension Case is required from the Superintending Engineer, PWD Electrical Circle, Udaipur.")
            else:
                notes.append("4. Approval of the Time Extension Case is to be done by this office.")
        else:
            notes.append("3. Work was completed in time.")
        if extra_item_amount > 0:
            extra_item_percentage = (extra_item_amount / float(work_order_amount) * 100) if work_order_amount > 0 else 0
            if extra_item_percentage > 5:
                notes.append(f"4. The amount of Extra items is Rs. {extra_item_amount} which is {extra_item_percentage:.2f}% of the Work Order Amount; exceed 5%, require approval from the Superintending Engineer, PWD Electrical Circle, Udaipur.")
            else:
                notes.append(f"4. The amount of Extra items is Rs. {extra_item_amount} which is {extra_item_percentage:.2f}% of the Work Order Amount; under 5%, approval of the same is to be granted by this office.")
        notes.extend([
            "5. Quality Control (QC) test reports attached.",
            "6. Please peruse above details for necessary decision-making.",
            "",
            "                                Premlata Jain",
            "                               AAO- As Auditor"
        ])

        note_sheet_data = {
            'agreement_no': work_order_data.get('agreement_no', '48/2024-25'),
            'name_of_work': work_order_data.get('name_of_work', 'Electric Repair and MTC work at Govt. Ambedkar hostel Ambamata, Govardhanvilas, Udaipur'),
            'name_of_firm': work_order_data.get('name_of_firm', 'M/s Seema Electrical Udaipur'),
            'date_commencement': work_order_data.get('date_commencement', '18/01/2025'),
            'date_completion': work_order_data.get('date_completion', '17/04/2025'),
            'actual_completion': work_order_data.get('actual_completion', '01/03/2025'),
            'work_order_amount': work_order_data.get('work_order_amount', '854678'),
            'extra_item_amount': extra_item_amount,
            'notes': notes,
            'totals': first_page_data.get('totals', {'payable': str(payable_amount)})
        }
        st.write(f"Note Sheet data: {note_sheet_data}")

        # Generate PDFs
        pdf_files = []
        for sheet_name, data, orientation in [
            ("First Page", first_page_data, "portrait"),
            ("Deviation Statement", deviation_data, "landscape"),
            ("Note Sheet", note_sheet_data, "portrait"),
        ]:
            pdf_path = os.path.join(TEMP_DIR, f"{sheet_name.replace(' ', '_')}.pdf")
            generate_pdf(sheet_name, data, orientation, pdf_path)
            pdf_files.append(pdf_path)

        # Merge PDFs
        current_date = datetime.now().strftime("%Y%m%d")
        pdf_output = os.path.join(TEMP_DIR, f"BILL_AND_DEVIATION_{current_date}.pdf")
        #############################################################################
        writer = PdfWriter()

        for pdf in pdf_files:
            if os.path.exists(pdf):
                reader = PdfReader(pdf)
                for page in reader.pages:
                    writer.add_page(page)

        with open(pdf_output, "wb") as out_file:
            writer.write(out_file)
        ###########################################################################

        # Generate Word docs
        word_files = []
        for sheet_name, data in [
            ("First Page", first_page_data),
            ("Last Page", last_page_data),
            ("Extra Items", extra_items_data),
            ("Deviation Statement", deviation_data),
            ("Note Sheet", note_sheet_data)
        ]:
            doc_path = os.path.join(TEMP_DIR, f"{sheet_name.replace(' ', '_')}.docx")
            create_word_doc(sheet_name, data, doc_path)
            word_files.append(doc_path)

        # Create ZIP
        zip_path = os.path.join(TEMP_DIR, "bill_output.zip")
        try:
            with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
                if os.path.exists(pdf_output):
                    zipf.write(pdf_output, os.path.basename(pdf_output))
                for word_file in word_files:
                    if os.path.exists(word_file):
                        zipf.write(word_file, os.path.basename(word_file))
            with open(zip_path, "rb") as f:
                st.download_button(
                    label="Download Bill Output",
                    data=f,
                    file_name="bill_output.zip",
                    mime="application/zip"
                )
        except Exception as e:
            st.error(f"Error creating ZIP file: {str(e)}")

        # Clean up temporary files
        import shutil
        try:
            shutil.rmtree(TEMP_DIR)
        except Exception as e:
            st.warning(f"Failed to clean up temp directory: {str(e)}")

    except Exception as e:
        st.error(f"Error: {str(e)}")
        st.write(traceback.format_exc())
