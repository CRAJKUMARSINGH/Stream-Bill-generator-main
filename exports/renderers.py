"""
Export renderers for generating PDF, Word, and other document formats
This module handles the output generation while preserving the core computation logic.

Foolproof PDF generation flow:
- Render HTML via Jinja2
- Generate PDF via a unified engine with intelligent fallbacks
- Optional in-memory caching to reduce repeated conversions
"""

import os
import tempfile
import json
import hashlib
from jinja2 import Environment, FileSystemLoader
from docx import Document
from pypdf import PdfReader, PdfWriter
import zipfile

# Unified PDF generator with fallbacks (weasyprint/reportlab/xhtml2pdf/pdfkit)
try:
    from core.pdf_generator_optimized import PDFGenerator
except Exception:  # Fallback for legacy path
    from pdf_generator_optimized import PDFGenerator  # type: ignore

# Lightweight in-memory cache (falls back silently if unavailable)
try:
    from data.cache_utils import get_cache
    _CACHE = get_cache()
except Exception:
    _CACHE = None


def setup_jinja_environment(template_dir):
    """Set up Jinja2 environment with the specified template directory"""
    # Keep no template cache to match existing test behavior
    return Environment(loader=FileSystemLoader(template_dir), cache_size=0)


def _hash_dict_stable(data: dict) -> str:
    """Create a stable hash for dictionaries (handles nested structures)."""
    try:
        payload = json.dumps(data, sort_keys=True, ensure_ascii=False).encode("utf-8")
    except Exception:
        payload = repr(data).encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def generate_html(sheet_name, data, template_dir, temp_dir):
    """
    Generate HTML file from template

    Args:
        sheet_name (str): Name of the sheet to generate
        data (dict): Data to render in the template
        template_dir (str): Directory containing templates
        temp_dir (str): Directory for temporary files

    Returns:
        str: Path to generated HTML file
    """
    env = setup_jinja_environment(template_dir)
    template = env.get_template(f"{sheet_name.lower().replace(' ', '_')}.html")
    html_content = template.render(data=data)
    html_path = os.path.join(temp_dir, f"{sheet_name.lower().replace(' ', '_')}.html")

    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)

    return html_path


def generate_pdf(sheet_name, data, orientation, template_dir, temp_dir, config=None):
    """
    Generate PDF via unified engine with robust fallbacks (no hard dependency on wkhtmltopdf).

    Args:
        sheet_name (str): Name of the sheet to generate
        data (dict): Data to render in the template
        orientation (str): Page orientation ("portrait" or "landscape")
        template_dir (str): Directory containing templates
        temp_dir (str): Directory for temporary files
        config: Unused; kept for backward compatibility

    Returns:
        str: Path to generated PDF file
    """
    env = setup_jinja_environment(template_dir)
    template = env.get_template(f"{sheet_name.lower().replace(' ', '_')}.html")
    html_content = template.render(data=data)

    os.makedirs(temp_dir, exist_ok=True)
    pdf_path = os.path.join(temp_dir, f"{sheet_name.replace(' ', '_')}.pdf")

    # Cache key for memoizing generated PDFs
    cache_key = None
    if _CACHE is not None:
        cache_key = f"pdf:{sheet_name}:{orientation}:{_hash_dict_stable(data)}"
        cached_path = _CACHE.get(cache_key)
        if isinstance(cached_path, str) and os.path.exists(cached_path):
            return cached_path

    # Note Sheet has special margins in the legacy flow; approximate in mm
    custom_margins = None
    if sheet_name == "Note Sheet":
        custom_margins = {"top": 6, "right": 6, "bottom": 15, "left": 6}

    generator = PDFGenerator(
        orientation=("landscape" if orientation == "landscape" else "portrait"),
        custom_margins=custom_margins,
    )

    success = generator.generate_pdf(html_content, pdf_path)
    if not success or not os.path.exists(pdf_path):
        raise RuntimeError("Failed to generate PDF with available engines")

    if _CACHE is not None and cache_key:
        try:
            _CACHE.set(cache_key, pdf_path, ttl=1800)  # 30 minutes
        except Exception:
            pass

    return pdf_path


def create_word_doc(sheet_name, data, doc_path):
    """
    Create Word document from data

    Args:
        sheet_name (str): Name of the sheet
        data (dict): Data to include in the document
        doc_path (str): Path where to save the document
    """
    doc = Document()
    if sheet_name == "First Page":
        table = doc.add_table(rows=len(data["items"]) + 3, cols=9)
        table.style = "Table Grid"
        for i, item in enumerate(data["items"]):
            row = table.rows[i]
            row.cells[0].text = str(item.get("unit", ""))
            row.cells[2].text = str(item.get("quantity", ""))
            row.cells[3].text = str(item.get("serial_no", ""))
            row.cells[4].text = str(item.get("description", ""))
            row.cells[5].text = str(item.get("rate", ""))
            row.cells[6].text = str(item.get("amount", ""))
            row.cells[8].text = str(item.get("remark", ""))
        row = table.rows[-3]
        row.cells[4].text = "Grand Total"
        row.cells[6].text = str(data["totals"].get("grand_total", ""))
        row = table.rows[-2]
        premium_percent_value = data['totals']['premium'].get('percent', 0)
        if isinstance(premium_percent_value, str):
            try:
                premium_percent_value = float(premium_percent_value)
            except (ValueError, TypeError):
                premium_percent_value = 0
        row.cells[4].text = f"Tender Premium @ {premium_percent_value:.2%}"
        row.cells[6].text = str(data["totals"]["premium"].get("amount", ""))
        row = table.rows[-1]
        row.cells[4].text = "Payable Amount"
        row.cells[6].text = str(data["totals"].get("payable", ""))
    elif sheet_name == "Last Page":
        doc.add_paragraph(f"Payable Amount: {data.get('payable_amount', '')}")
        doc.add_paragraph(f"Total in Words: {data.get('amount_words', '')}")
    elif sheet_name == "Extra Items":
        table = doc.add_table(rows=len(data["items"]) + 1, cols=7)
        table.style = "Table Grid"
        headers = ["Serial No.", "Remark", "Description", "Quantity", "Unit", "Rate", "Amount"]
        for j, header in enumerate(headers):
            table.rows[0].cells[j].text = header
        for i, item in enumerate(data["items"]):
            row = table.rows[i + 1]
            row.cells[0].text = str(item.get("serial_no", ""))
            row.cells[1].text = str(item.get("remark", ""))
            row.cells[2].text = str(item.get("description", ""))
            row.cells[3].text = str(item.get("quantity", ""))
            row.cells[4].text = str(item.get("unit", ""))
            row.cells[5].text = str(item.get("rate", ""))
            row.cells[6].text = str(item.get("amount", ""))
    elif sheet_name == "Deviation Statement":
        table = doc.add_table(rows=len(data["items"]) + 5, cols=12)
        table.style = "Table Grid"
        headers = ["Serial No.", "Description", "Unit", "Qty WO", "Rate", "Amt WO", "Qty Bill", "Amt Bill", "Excess Qty", "Excess Amt", "Saving Qty", "Saving Amt"]
        for j, header in enumerate(headers):
            table.rows[0].cells[j].text = header
        for i, item in enumerate(data["items"]):
            row = table.rows[i + 1]
            row.cells[0].text = str(item.get("serial_no", ""))
            row.cells[1].text = str(item.get("description", ""))
            row.cells[2].text = str(item.get("unit", ""))
            row.cells[3].text = str(item.get("qty_wo", ""))
            row.cells[4].text = str(item.get("rate", ""))
            row.cells[5].text = str(item.get("amt_wo", ""))
            row.cells[6].text = str(item.get("qty_bill", ""))
            row.cells[7].text = str(item.get("amt_bill", ""))
            row.cells[8].text = str(item.get("excess_qty", ""))
            row.cells[9].text = str(item.get("excess_amt", ""))
            row.cells[10].text = str(item.get("saving_qty", ""))
            row.cells[11].text = str(item.get("saving_amt", ""))
        row = table.rows[-4]
        row.cells[1].text = "Grand Total"
        row.cells[5].text = str(data["summary"].get("work_order_total", ""))
        row.cells[7].text = str(data["summary"].get("executed_total", ""))
        row.cells[9].text = str(data["summary"].get("overall_excess", ""))
        row.cells[11].text = str(data["summary"].get("overall_saving", ""))
        row = table.rows[-3]
        premium_percent_value = data['summary']['premium'].get('percent', 0)
        if isinstance(premium_percent_value, str):
            try:
                premium_percent_value = float(premium_percent_value)
            except (ValueError, TypeError):
                premium_percent_value = 0
        row.cells[1].text = f"Add Tender Premium ({premium_percent_value:.2%})"
        row.cells[5].text = str(data["summary"].get("tender_premium_f", ""))
        row.cells[7].text = str(data["summary"].get("tender_premium_h", ""))
        row.cells[9].text = str(data["summary"].get("tender_premium_j", ""))
        row.cells[11].text = str(data["summary"].get("tender_premium_l", ""))
        row = table.rows[-2]
        row.cells[1].text = "Grand Total including Tender Premium"
        row.cells[5].text = str(data["summary"].get("grand_total_f", ""))
        row.cells[7].text = str(data["summary"].get("grand_total_h", ""))
        row.cells[9].text = str(data["summary"].get("grand_total_j", ""))
        row.cells[11].text = str(data["summary"].get("grand_total_l", ""))
        row = table.rows[-1]
        net_difference = data["summary"].get("net_difference", 0)
        row.cells[1].text = "Overall Excess" if net_difference > 0 else "Overall Saving"
        row.cells[7].text = str(abs(round(net_difference)))
    elif sheet_name == "Note Sheet":
        for note in data.get("notes", []):
            doc.add_paragraph(str(note))
    elif sheet_name == "Certificate II":
        # Add certificate II content
        doc.add_heading("II. CERTIFICATE AND SIGNATURES", level=1)

        # Certificate text
        certificate_text = f"""The measurements on which are based the entries in columns 1 to 6 of Account I, were made by {data.get('measurement_officer', 'Junior Engineer')} on {data.get('measurement_date', '01/03/2025')}, and are recorded at page {data.get('measurement_book_page', '04-20')} of Measurement Book No. {data.get('measurement_book_no', '887')}."""
        doc.add_paragraph(certificate_text)

        doc.add_paragraph("*Certified that in addition to and quite apart from the quantities of work actually executed, as shown in column 4 of Account I, some work has actually been done in connection with several items and the value of such work (after deduction therefrom the proportionate amount of secured advances, if any, ultimately recoverable on account of the quantities of materials used therein) is in no case, less than the advance payments as per item 2 of the Memorandum, if payment is made.")

        doc.add_paragraph("+Certified that the contractor has made satisfactory progress with the work, and that the quantities and amounts claimed are correct and the work has been executed in accordance with the specifications and the terms of the contract.")

        doc.add_paragraph("I also certify that the amount claimed is not more than the amount admissible under the contract.")

        # Signature blocks
        doc.add_paragraph("\nDated signature of officer preparing the bill")
        doc.add_paragraph(f"Name: {data.get('officer_name', 'Name of Officer')}")
        doc.add_paragraph(f"Designation: {data.get('officer_designation', 'Assistant Engineer')}")
        doc.add_paragraph(f"Date: {data.get('bill_date', '__/__/____')}")

        doc.add_paragraph("\n+Dated signature of officer authorising payment")
        doc.add_paragraph(f"Name: {data.get('authorising_officer_name', 'Name of Authorising Officer')}")
        doc.add_paragraph(f"Designation: {data.get('authorising_officer_designation', 'Executive Engineer')}")
        doc.add_paragraph(f"Date: {data.get('authorisation_date', '__/__/____')}")

    elif sheet_name == "Certificate III":
        # Add certificate III content
        doc.add_heading("III. MEMORANDUM OF PAYMENTS", level=1)

        # Create table for payment details
        table = doc.add_table(rows=25, cols=4)
        table.style = "Table Grid"

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "S.No."
        hdr_cells[1].text = "Description"
        hdr_cells[2].text = "Entry No."
        hdr_cells[3].text = "Amount Rs."

        # Data rows
        table.rows[1].cells[0].text = "1."
        table.rows[1].cells[1].text = "Total value of work actually measured, as per Account I, Col. 5, Entry [A]"
        table.rows[1].cells[2].text = "[A]"
        table.rows[1].cells[3].text = str(data["totals"].get("grand_total", "0"))

        table.rows[2].cells[0].text = "2."
        table.rows[2].cells[1].text = "Total up-to-date advance payments for work not yet measured as per details given below:"
        table.rows[2].cells[2].text = ""
        table.rows[2].cells[3].text = ""

        table.rows[3].cells[0].text = ""
        table.rows[3].cells[1].text = "(a) Total as per previous bill"
        table.rows[3].cells[2].text = "[B]"
        table.rows[3].cells[3].text = "Nil"

        table.rows[4].cells[0].text = ""
        table.rows[4].cells[1].text = "(b) Since previous bill"
        table.rows[4].cells[2].text = "[D]"
        table.rows[4].cells[3].text = "Nil"

        table.rows[5].cells[0].text = "3."
        table.rows[5].cells[1].text = "Total up-to-date secured advances on security of materials"
        table.rows[5].cells[2].text = "[C]"
        table.rows[5].cells[3].text = "Nil"

        table.rows[6].cells[0].text = "4."
        table.rows[6].cells[1].text = "Total (Items 1 + 2 + 3) A+B+C"
        table.rows[6].cells[2].text = ""
        table.rows[6].cells[3].text = str(data["totals"].get("grand_total", "0"))

        table.rows[7].cells[0].text = "5."
        table.rows[7].cells[1].text = "Deduct: Amount withheld"
        table.rows[7].cells[2].text = ""
        table.rows[7].cells[3].text = ""

        table.rows[8].cells[0].text = ""
        table.rows[8].cells[1].text = "(a) From previous bill as per last Running Account Bill"
        table.rows[8].cells[2].text = "[5]"
        table.rows[8].cells[3].text = "Nil"

        table.rows[9].cells[0].text = ""
        table.rows[9].cells[1].text = "(b) From this bill"
        table.rows[9].cells[2].text = ""
        table.rows[9].cells[3].text = "Nil"

        table.rows[10].cells[0].text = "6."
        table.rows[10].cells[1].text = 'Balance i.e. "up-to-date" payments (Item 4-5)'
        table.rows[10].cells[2].text = ""
        table.rows[10].cells[3].text = str(data["totals"].get("grand_total", "0"))

        table.rows[11].cells[0].text = "7."
        table.rows[11].cells[1].text = "Total amount of payments already made as per Entry (K)"
        table.rows[11].cells[2].text = "[K]"
        table.rows[11].cells[3].text = "0"

        table.rows[12].cells[0].text = "8."
        table.rows[12].cells[1].text = "Payments now to be made, as detailed below:"
        table.rows[12].cells[2].text = ""
        table.rows[12].cells[3].text = str(data["totals"].get("payable", "0"))

        table.rows[13].cells[0].text = ""
        table.rows[13].cells[1].text = "(a) By recovery of amounts creditable to this work"
        table.rows[13].cells[2].text = "[a]"
        table.rows[13].cells[3].text = ""

        # Calculate deductions
        payable_amount = float(data["totals"].get("payable", 0))
        sd_amount = payable_amount * 0.10
        it_amount = payable_amount * 0.02
        gst_amount = payable_amount * 0.02
        lc_amount = payable_amount * 0.01

        table.rows[14].cells[0].text = ""
        table.rows[14].cells[1].text = "SD @ 10%"
        table.rows[14].cells[2].text = ""
        table.rows[14].cells[3].text = f"{sd_amount:.0f}"

        table.rows[15].cells[0].text = ""
        table.rows[15].cells[1].text = "IT @ 2%"
        table.rows[15].cells[2].text = ""
        table.rows[15].cells[3].text = f"{it_amount:.0f}"

        table.rows[16].cells[0].text = ""
        table.rows[16].cells[1].text = "GST @ 2%"
        table.rows[16].cells[2].text = ""
        table.rows[16].cells[3].text = f"{gst_amount:.0f}"

        table.rows[17].cells[0].text = ""
        table.rows[17].cells[1].text = "LC @ 1%"
        table.rows[17].cells[2].text = ""
        table.rows[17].cells[3].text = f"{lc_amount:.0f}"

        total_deductions = sd_amount + it_amount + gst_amount + lc_amount
        table.rows[18].cells[0].text = ""
        table.rows[18].cells[1].text = "Total recovery"
        table.rows[18].cells[2].text = ""
        table.rows[18].cells[3].text = f"{total_deductions:.0f}"

        table.rows[19].cells[0].text = ""
        table.rows[19].cells[1].text = "(b) By recovery of amount creditable to other works"
        table.rows[19].cells[2].text = "[b]"
        table.rows[19].cells[3].text = "Nil"

        cheque_amount = payable_amount - total_deductions
        table.rows[20].cells[0].text = ""
        table.rows[20].cells[1].text = "(c) By cheque"
        table.rows[20].cells[2].text = "[c]"
        table.rows[20].cells[3].text = f"{cheque_amount:.0f}"

        # Payment details
        doc.add_paragraph(f"\nPay Rs. {cheque_amount:.0f}")
        doc.add_paragraph(f"Pay Rupees {data.get('payable_words', 'Zero')} (by cheque)")
        doc.add_paragraph("Dated the ____ / ____ / ________")
        doc.add_paragraph("Dated initials of Disbursing Officer: _______________")
        doc.add_paragraph(f"\nReceived Rupees {data.get('payable_words', 'Zero')} (by cheque) as per above memorandum, on account of this bill")
        doc.add_paragraph("Signature of Contractor: _______________")
        doc.add_paragraph("\nPaid by me, vide cheque No. _______ dated ____ / ____ / ________")
        doc.add_paragraph("Dated initials of person actually making the payment: _______________")

    doc.save(doc_path)


def merge_pdfs(pdf_files, output_path):
    """
    Merge multiple PDF files into a single PDF

    Args:
        pdf_files (list): List of paths to PDF files to merge
        output_path (str): Path where to save the merged PDF
    """
    writer = PdfWriter()

    for pdf in pdf_files:
        if os.path.exists(pdf):
            reader = PdfReader(pdf)
            for page in reader.pages:
                writer.add_page(page)

    with open(output_path, "wb") as out_file:
        writer.write(out_file)


def create_zip_archive(files, zip_path):
    """
    Create a ZIP archive containing the specified files

    Args:
        files (list): List of file paths to include in the ZIP
        zip_path (str): Path where to save the ZIP archive
    """
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
        for file_path in files:
            if os.path.exists(file_path):
                zipf.write(file_path, os.path.basename(file_path))
